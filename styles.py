# -*- coding: utf-8 -*-
"""
Word样式定义模块 - LaTeX风格样式
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# 字体配置
FONTS = {
    'heading_cn': '黑体',            # 标题中文字体
    'body_cn': '宋体',               # 正文中文字体
    'body_en': 'Times New Roman',    # 正文英文字体
    'mono': 'Consolas',              # 代码字体
    'math': 'Cambria Math',          # 数学字体
}


def _pt(value, fallback: Pt) -> Pt:
    try:
        if value is None:
            return fallback
        if isinstance(value, Pt):
            return value
        return Pt(float(value))
    except Exception:
        return fallback


def _cm(value, fallback: Cm) -> Cm:
    try:
        if value is None:
            return fallback
        if isinstance(value, Cm):
            return value
        return Cm(float(value))
    except Exception:
        return fallback


def _align(value, fallback=WD_ALIGN_PARAGRAPH.LEFT):
    try:
        v = str(value or '').strip().lower()
        if v in ('center', 'centre'):
            return WD_ALIGN_PARAGRAPH.CENTER
        if v in ('right',):
            return WD_ALIGN_PARAGRAPH.RIGHT
        if v in ('justify', 'both'):
            return WD_ALIGN_PARAGRAPH.JUSTIFY
        return WD_ALIGN_PARAGRAPH.LEFT
    except Exception:
        return fallback

# 中国标准字号对应磅值
FONT_SIZES = {
    'er_hao': Pt(22),       # 二号 - 一级标题
    'san_hao': Pt(16),      # 三号 - 二级标题
    'xiao_san': Pt(15),     # 小三 - 三级标题
    'si_hao': Pt(14),       # 四号 - 四级标题
    'xiao_si': Pt(12),      # 小四 - 正文
    'wu_hao': Pt(10.5),     # 五号
    'code': Pt(10),         # 代码
    'caption': Pt(10.5),    # 图表标题
}

# 兼容旧代码的别名
LATEX_FONTS = FONTS

# 样式配置
STYLE_CONFIGS = {
    'standard': {
        'heading_cn': '黑体',
        'body_cn': '宋体',
        'body_en': 'Times New Roman',
        'body_size': Pt(12),  # 小四
    },
    'academic': {
        'heading_cn': '黑体',
        'body_cn': '宋体',
        'body_en': 'Times New Roman',
        'body_size': Pt(12),  # 小四
    },
    'simple': {
        'heading_cn': '微软雅黑',
        'body_cn': '微软雅黑',
        'body_en': 'Arial',
        'body_size': Pt(11),
    }
}


def setup_document_styles(doc: Document, style: str = "standard", page_size: str = "a4", style_overrides: dict = None) -> None:
    """设置文档的样式
    
    Args:
        doc: Word文档对象
        style: 样式类型 (standard/academic/simple)
        page_size: 页面大小 (a4/letter)
    """
    style_config = STYLE_CONFIGS.get(style, STYLE_CONFIGS['standard'])
    if isinstance(style_overrides, dict) and style_overrides:
        style_config = {**style_config, **style_overrides}

    # 保存到 doc 上，供页边距等函数读取
    try:
        doc._style_config = style_config
    except Exception:
        pass
    
    # 设置默认字体
    _set_default_font(doc, style_config)
    
    # 创建/修改标题样式
    _setup_heading_styles(doc, style_config)
    
    # 创建正文样式
    _setup_body_styles(doc, style_config)
    
    # 创建代码块样式
    _setup_code_styles(doc, style_config)
    
    # 创建引用样式
    _setup_quote_styles(doc, style_config)
    
    # 设置页面边距和大小
    _setup_page_margins(doc, page_size)


def _set_default_font(doc: Document, style_config: dict = None) -> None:
    """设置文档默认字体"""
    if style_config is None:
        style_config = STYLE_CONFIGS['standard']
    
    style = doc.styles['Normal']
    font = style.font
    font.name = style_config.get('body_en', FONTS['body_en'])
    font.size = _pt(style_config.get('body_size_pt'), style_config.get('body_size', FONT_SIZES['xiao_si']))
    font.bold = False
    
    # 设置中文字体
    r = style.element.rPr
    if r is None:
        r = OxmlElement('w:rPr')
        style.element.append(r)
    
    rFonts = r.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        r.append(rFonts)
    rFonts.set(qn('w:eastAsia'), style_config.get('body_cn', FONTS['body_cn']))


def _setup_heading_styles(doc: Document, style_config: dict = None) -> None:
    """设置标题样式"""
    if style_config is None:
        style_config = STYLE_CONFIGS['standard']
    
    heading_cn = style_config.get('heading_cn', FONTS['heading_cn'])
    body_en = style_config.get('body_en', FONTS['body_en'])
    
    # Heading 1 - 一级标题：居中、黑体、加粗、二号
    h1 = doc.styles['Heading 1']
    h1.font.name = body_en
    h1.font.size = _pt(style_config.get('heading1_size_pt'), FONT_SIZES['er_hao'])
    h1.font.bold = bool(style_config.get('heading1_bold', True))
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.alignment = _align(style_config.get('heading1_alignment'), WD_ALIGN_PARAGRAPH.CENTER)
    h1.paragraph_format.space_before = _pt(style_config.get('heading1_space_before_pt'), Pt(24))
    h1.paragraph_format.space_after = _pt(style_config.get('heading1_space_after_pt'), Pt(18))
    h1.paragraph_format.keep_with_next = True
    _set_heading_cjk_font(h1, heading_cn)
    
    # Heading 2 - 二级标题：居中、三号、加粗、黑体
    h2 = doc.styles['Heading 2']
    h2.font.name = body_en
    h2.font.size = _pt(style_config.get('heading2_size_pt'), FONT_SIZES['san_hao'])
    h2.font.bold = bool(style_config.get('heading2_bold', True))
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.alignment = _align(style_config.get('heading2_alignment'), WD_ALIGN_PARAGRAPH.CENTER)
    h2.paragraph_format.space_before = _pt(style_config.get('heading2_space_before_pt'), Pt(18))
    h2.paragraph_format.space_after = _pt(style_config.get('heading2_space_after_pt'), Pt(12))
    h2.paragraph_format.keep_with_next = True
    _set_heading_cjk_font(h2, heading_cn)
    
    # Heading 3 - 三级标题：靠左、小三、加粗、黑体
    h3 = doc.styles['Heading 3']
    h3.font.name = body_en
    h3.font.size = _pt(style_config.get('heading3_size_pt'), FONT_SIZES['xiao_san'])
    h3.font.bold = bool(style_config.get('heading3_bold', True))
    h3.font.italic = False
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3.paragraph_format.alignment = _align(style_config.get('heading3_alignment'), WD_ALIGN_PARAGRAPH.LEFT)
    h3.paragraph_format.space_before = _pt(style_config.get('heading3_space_before_pt'), Pt(13))
    h3.paragraph_format.space_after = _pt(style_config.get('heading3_space_after_pt'), Pt(10))
    h3.paragraph_format.keep_with_next = True
    _set_heading_cjk_font(h3, heading_cn)
    
    # Heading 4 - 四级标题：靠左、四号、加粗、黑体
    h4 = doc.styles['Heading 4']
    h4.font.name = body_en
    h4.font.size = _pt(style_config.get('heading4_size_pt'), FONT_SIZES['si_hao'])
    h4.font.bold = bool(style_config.get('heading4_bold', True))
    h4.font.italic = False
    h4.font.color.rgb = RGBColor(0, 0, 0)
    h4.paragraph_format.alignment = _align(style_config.get('heading4_alignment'), WD_ALIGN_PARAGRAPH.LEFT)
    h4.paragraph_format.space_before = _pt(style_config.get('heading4_space_before_pt'), Pt(10))
    h4.paragraph_format.space_after = _pt(style_config.get('heading4_space_after_pt'), Pt(6))
    h4.paragraph_format.keep_with_next = True
    _set_heading_cjk_font(h4, heading_cn)


def _set_heading_cjk_font(style, font_name: str = '黑体') -> None:
    """为标题设置中文字体"""
    r = style.element.rPr
    if r is None:
        r = OxmlElement('w:rPr')
        style.element.append(r)
    rFonts = r.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        r.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def _setup_body_styles(doc: Document, style_config: dict = None) -> None:
    """设置正文样式"""
    if style_config is None:
        style_config = STYLE_CONFIGS['standard']
    normal = doc.styles['Normal']
    normal.paragraph_format.alignment = _align(style_config.get('body_alignment'), WD_ALIGN_PARAGRAPH.LEFT)

    body_ls = style_config.get('body_line_spacing', 1.5)
    try:
        body_ls_f = float(body_ls)
    except Exception:
        body_ls_f = 1.5
    if abs(body_ls_f - 1.0) < 0.01:
        normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    elif abs(body_ls_f - 1.5) < 0.01:
        normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    elif abs(body_ls_f - 2.0) < 0.01:
        normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    else:
        normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        normal.paragraph_format.line_spacing = body_ls_f

    normal.paragraph_format.space_before = _pt(style_config.get('body_space_before_pt'), Pt(0))
    normal.paragraph_format.space_after = _pt(style_config.get('body_space_after_pt'), Pt(6))
    normal.paragraph_format.first_line_indent = _pt(style_config.get('body_first_line_indent_pt'), Pt(24))


def _setup_code_styles(doc: Document, style_config: dict = None) -> None:
    """创建代码块样式"""
    styles = doc.styles
    
    # 检查样式是否存在
    if 'Code Block' not in [s.name for s in styles]:
        code_style = styles.add_style('Code Block', WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = styles['Code Block']
    
    if style_config is None:
        style_config = STYLE_CONFIGS['standard']
    code_style.font.name = style_config.get('mono', FONTS['mono'])
    code_style.font.size = _pt(style_config.get('code_size_pt'), FONT_SIZES['code'])
    code_style.paragraph_format.space_before = _pt(style_config.get('code_space_before_pt'), Pt(6))
    code_style.paragraph_format.space_after = _pt(style_config.get('code_space_after_pt'), Pt(6))
    code_style.paragraph_format.left_indent = _cm(style_config.get('code_left_indent_cm'), Cm(0.5))
    try:
        code_style.paragraph_format.line_spacing = float(style_config.get('code_line_spacing', 1.0))
    except Exception:
        code_style.paragraph_format.line_spacing = 1.0
    
    # 行内代码样式
    if 'Inline Code' not in [s.name for s in styles]:
        inline_code = styles.add_style('Inline Code', WD_STYLE_TYPE.CHARACTER)
    else:
        inline_code = styles['Inline Code']
    
    inline_code.font.name = style_config.get('mono', FONTS['mono'])
    inline_code.font.size = _pt(style_config.get('code_size_pt'), FONT_SIZES['code'])


def _setup_quote_styles(doc: Document, style_config: dict = None) -> None:
    """创建引用块样式"""
    if style_config is None:
        style_config = STYLE_CONFIGS['standard']
    styles = doc.styles
    
    if 'Block Quote' not in [s.name for s in styles]:
        quote_style = styles.add_style('Block Quote', WD_STYLE_TYPE.PARAGRAPH)
    else:
        quote_style = styles['Block Quote']
    
    quote_style.font.name = style_config.get('quote_font', FONTS['body_en'])
    quote_style.font.size = _pt(style_config.get('quote_size_pt'), FONT_SIZES['xiao_si'])
    quote_style.font.italic = bool(style_config.get('quote_italic', True))
    quote_style.paragraph_format.left_indent = _cm(style_config.get('quote_left_indent_cm'), Cm(1))
    quote_style.paragraph_format.right_indent = _cm(style_config.get('quote_right_indent_cm'), Cm(1))
    quote_style.paragraph_format.space_before = _pt(style_config.get('quote_space_before_pt'), Pt(6))
    quote_style.paragraph_format.space_after = _pt(style_config.get('quote_space_after_pt'), Pt(6))


def _setup_page_margins(doc: Document, page_size: str = "a4") -> None:
    """设置页面边距和大小"""
    for section in doc.sections:
        # 设置页面大小
        if page_size == "letter":
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
        else:  # a4
            section.page_width = Cm(21)
            section.page_height = Cm(29.7)
        
        # 设置边距
        # 允许覆盖页边距
        style_config = getattr(doc, '_style_config', None)
        if isinstance(style_config, dict):
            section.top_margin = _cm(style_config.get('margin_top_cm'), Cm(2.54))
            section.bottom_margin = _cm(style_config.get('margin_bottom_cm'), Cm(2.54))
            section.left_margin = _cm(style_config.get('margin_left_cm'), Cm(3.18))
            section.right_margin = _cm(style_config.get('margin_right_cm'), Cm(3.18))
        else:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.18)
            section.right_margin = Cm(3.18)


def apply_table_style(table, style_config: dict = None) -> None:
    """应用LaTeX风格的表格样式 - 三线表"""
    if isinstance(style_config, dict):
        align = str(style_config.get('table_alignment', 'center')).lower()
        if align == 'left':
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
        elif align == 'right':
            table.alignment = WD_TABLE_ALIGNMENT.RIGHT
        else:
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
    else:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 移除所有边框
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    
    # 设置表格边框
    tblBorders = OxmlElement('w:tblBorders')
    
    # 只设置顶部、底部边框和表头下边框（三线表）
    for border_name in ['top', 'bottom']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')  # 1.5pt
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    
    # 清除左右和内部边框
    for border_name in ['left', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tblBorders.append(border)
    
    tblPr.append(tblBorders)
    
    # 设置表头行下边框
    if len(table.rows) > 0:
        header_row = table.rows[0]
        for cell in header_row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')  # 0.75pt
            bottom.set(qn('w:color'), '000000')
            tcBorders.append(bottom)
            tcPr.append(tcBorders)
            
            header_bold = True
            if isinstance(style_config, dict):
                header_bold = bool(style_config.get('table_header_bold', True))
            if header_bold:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True


def create_caption_paragraph(doc: Document, text: str, prefix: str = "图") -> None:
    """创建图表标题段落"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = para.add_run(f"{prefix} {text}")
    run.font.size = FONT_SIZES['caption']
    run.font.name = FONTS['body_en']
    
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(12)
