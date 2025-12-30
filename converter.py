# -*- coding: utf-8 -*-
"""
Markdown转Word核心转换器
"""

import re
import os
import logging
logging.getLogger('matplotlib').setLevel(logging.ERROR)

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from styles import setup_document_styles, FONT_SIZES, FONTS
from handlers import ImageHandler, TableHandler, MathHandler, CodeHandler, ListHandler
from utils import clean_text, is_valid_math_formula, normalize_markdown, convert_latex_delimiters
from parser import parse_markdown, parse_inline, parse_table, InlineType, BlockElement


class ExportCancelled(Exception):
    pass


class ConversionError(Exception):
    pass


class MarkdownToWordConverter:
    """Markdown转Word转换器主类"""
    
    def __init__(self, base_dir: str = None, style: str = "standard", page_size: str = "a4", export_style: dict = None, template_path: str = None, header_footer_config: dict = None):
        """
        初始化转换器
        
        Args:
            base_dir: Markdown文件所在目录，用于解析相对路径的图片
            style: 文档样式 (standard/academic/simple)
            page_size: 页面大小 (a4/letter)
            template_path: Word模板路径 (.docx)
        """
        self.base_dir = base_dir
        self.style = style
        self.page_size = page_size
        self.template_path = template_path
        self.header_footer_config = header_footer_config
        self.doc = None

        self.export_style = export_style if isinstance(export_style, dict) else {}

        # 诊断信息（用于导出报告）
        self.diagnostics = {
            'unresolved_refs': [],
            'include_issues': [],
            'included_files': [],
        }

        # 交叉引用：id -> bookmark / label
        self._bookmark_id_counter = 1
        self._id_to_bookmark = {}
        self._id_to_label = {}
        
        # 脚注存储
        self.footnotes = {}
        self._fig_counter = 0
        self._tbl_counter = 0
        self._sec_counters = [0, 0, 0, 0]
        
        # 初始化各处理器
        self.image_handler = ImageHandler(base_dir, style_config=self.export_style)
        self.table_handler = TableHandler(style_config=self.export_style)
        self.math_handler = MathHandler()
        self.code_handler = CodeHandler()
        self.list_handler = ListHandler()
        
        # 有序列表编号计数器，用于创建独立编号定义
        # 从1000开始避免与文档已有的numId冲突
        self._num_id_counter = 1000
        self._current_num_id = 1000  # 当前列表的编号ID
    
    def convert_file(self, input_path: str, output_path: str) -> bool:
        """
        转换Markdown文件为Word文档
        
        Args:
            input_path: Markdown文件路径
            output_path: 输出Word文件路径
            
        Returns:
            是否成功
        """
        try:
            # 更新base_dir
            if self.base_dir is None:
                self.base_dir = os.path.dirname(os.path.abspath(input_path))
                self.image_handler.base_dir = self.base_dir
            
            # 读取Markdown内容
            with open(input_path, 'r', encoding='utf-8') as f:
                markdown_text = f.read()
            
            # 转换
            self.convert_text(markdown_text)
            
            # 保存
            self.doc.save(output_path)
            print(f"转换成功: {output_path}")
            return True
            
        except Exception as e:
            print(f"转换失败: {e}")
            return False
    
    def convert_text(self, markdown_text: str, progress_callback=None, cancel_event=None) -> Document:
        """
        转换Markdown文本为Word文档对象
        
        Args:
            markdown_text: Markdown格式文本
            
        Returns:
            Document对象
        """
        # 创建新文档
        if self.template_path and os.path.exists(self.template_path):
            try:
                self.doc = Document(self.template_path)
            except Exception:
                self.doc = Document()
        else:
            self.doc = Document()
        setup_document_styles(self.doc, style=self.style, page_size=self.page_size, style_overrides=self.export_style)

        # 应用页眉页脚配置
        if self.header_footer_config:
            self._apply_header_footer()
            
        # 是否打开时自动更新字段（目录/编号等）
        try:
            self._set_update_fields_on_open(bool(self.export_style.get('update_fields_on_open', True)))
        except Exception:
            pass
        
        # 预处理文本
        try:
            markdown_text = self._expand_includes(markdown_text)
        except Exception:
            pass

        markdown_text = clean_text(markdown_text)
        markdown_text = convert_latex_delimiters(markdown_text)  # 转换 \(...\) 和 \[...\] 格式
        markdown_text = normalize_markdown(markdown_text)  # 规范化格式
        
        # 分割成块并处理
        blocks = self._split_into_blocks(markdown_text)
        try:
            self._prescan_blocks_for_refs(blocks)
        except Exception:
            pass
        total = len(blocks)

        for idx, blk in enumerate(blocks, start=1):
            if cancel_event is not None and getattr(cancel_event, 'is_set', lambda: False)():
                raise ExportCancelled('导出已取消')

            block_type = blk.get('type')
            content = blk.get('content')
            start_line = int(blk.get('start_line') or 1)

            if callable(progress_callback):
                try:
                    progress_callback(idx, total, block_type, start_line)
                except Exception:
                    pass

            try:
                self._process_block(block_type, content)
            except ExportCancelled:
                raise
            except Exception as e:
                snippet = ''
                try:
                    if isinstance(content, str):
                        snippet = content.strip().replace('\n', ' ')[:180]
                    else:
                        snippet = str(content)[:180]
                except Exception:
                    snippet = ''
                raise ConversionError(
                    f"在第{start_line}行附近处理 {block_type} 失败: {e}" + (f"\n内容片段: {snippet}" if snippet else "")
                ) from e
        # 原生脚注已在 _add_inline_content 中通过 _add_native_footnote 处理
        # 无需额外渲染步骤
            
        return self.doc

    def _prescan_blocks_for_refs(self, blocks: list) -> None:
        """预扫描 blocks：为 {#id} 分配书签名与显示标签，支持前向引用。
        同时预先收集所有脚注定义，以便在后续处理中使用。"""
        try:
            self._fig_counter = 0
            self._tbl_counter = 0
            self._sec_counters = [0, 0, 0, 0]
        except Exception:
            pass

        for blk in (blocks or []):
            try:
                t = blk.get('type')
                c = blk.get('content')
                anchor = None
                level = None

                # 预先收集脚注定义
                if t == 'footnote_def':
                    if isinstance(c, dict):
                        ref = c.get('ref')
                        text = c.get('text')
                        if ref:
                            self.footnotes[ref] = text or ''
                    continue

                if t and t.startswith('heading_'):
                    level = int(t.split('_')[1])
                    if isinstance(c, dict):
                        anchor = c.get('anchor')
                elif t == 'image':
                    if isinstance(c, dict):
                        anchor = c.get('anchor')
                elif t == 'table':
                    if isinstance(c, dict):
                        anchor = c.get('anchor')

                if anchor:
                    self._ensure_bookmark(anchor)
                    if str(anchor).startswith('sec:') and level is not None:
                        self._make_label_for_anchor(anchor, 'heading', heading_level=level)
                    elif str(anchor).startswith('fig:'):
                        self._make_label_for_anchor(anchor, 'image')
                    elif str(anchor).startswith('tbl:'):
                        self._make_label_for_anchor(anchor, 'table')
            except Exception:
                pass

    def _expand_includes(self, text: str) -> str:
        """展开 @include 指令（支持引号、相对路径、递归与循环保护）。"""
        base_dir = self.base_dir
        if not base_dir:
            try:
                base_dir = os.getcwd()
            except Exception:
                base_dir = None
        return self._expand_includes_inner(text, current_dir=base_dir, stack=[])

    def _expand_includes_inner(self, text: str, current_dir: str, stack: list) -> str:
        out_lines = []
        lines = (text or '').split('\n')
        pat = re.compile(r'^\s*@include\s+("[^"]+"|\S+)\s*$')
        for ln in lines:
            m = pat.match(ln or '')
            if not m:
                out_lines.append(ln)
                continue

            raw = m.group(1)
            p = raw.strip()
            if p.startswith('"') and p.endswith('"') and len(p) >= 2:
                p = p[1:-1]
            p = p.strip()
            if not p:
                out_lines.append(ln)
                continue

            inc_path = p
            try:
                if current_dir and not os.path.isabs(inc_path):
                    inc_path = os.path.normpath(os.path.join(current_dir, inc_path))
                inc_path = os.path.abspath(inc_path)
            except Exception:
                pass

            # 循环保护
            if inc_path in stack:
                out_lines.append(f"[include循环: {p}]")
                try:
                    self.diagnostics['include_issues'].append({'type': 'cycle', 'path': inc_path})
                except Exception:
                    pass
                continue

            if not os.path.exists(inc_path):
                out_lines.append(f"[include缺失: {p}]")
                try:
                    self.diagnostics['include_issues'].append({'type': 'missing', 'path': inc_path})
                except Exception:
                    pass
                continue

            try:
                with open(inc_path, 'r', encoding='utf-8') as f:
                    inc_text = f.read()
            except Exception:
                out_lines.append(f"[include读取失败: {p}]")
                try:
                    self.diagnostics['include_issues'].append({'type': 'read_error', 'path': inc_path})
                except Exception:
                    pass
                continue

            try:
                self.diagnostics['included_files'].append(inc_path)
            except Exception:
                pass

            next_dir = None
            try:
                next_dir = os.path.dirname(inc_path)
            except Exception:
                next_dir = current_dir

            expanded = self._expand_includes_inner(inc_text, current_dir=next_dir, stack=stack + [inc_path])
            out_lines.append(expanded)

        return '\n'.join(out_lines)

    def _use_word_styles(self) -> bool:
        try:
            return bool(self.export_style.get('use_word_styles', False))
        except Exception:
            return False

    def _mapped_style_name(self, key: str, fallback: str = None) -> str:
        try:
            v = self.export_style.get(key)
            if v is None:
                return fallback
            s = str(v).strip()
            return s if s else fallback
        except Exception:
            return fallback

    def _apply_paragraph_style(self, para, style_name: str) -> None:
        if not style_name:
            return
        try:
            # style_name 不存在时会抛 KeyError
            para.style = style_name
        except Exception:
            pass

    def _set_update_fields_on_open(self, enabled: bool) -> None:
        """设置 Word 打开文档时是否自动更新域（fields）。"""
        try:
            settings = self.doc.settings.element
            # 先移除已有的 updateFields
            try:
                for el in settings.findall(qn('w:updateFields')):
                    settings.remove(el)
            except Exception:
                pass

            el = OxmlElement('w:updateFields')
            el.set(qn('w:val'), 'true' if enabled else 'false')
            settings.append(el)
        except Exception:
            pass
    
    def _split_into_blocks(self, text: str) -> list:
        """
        将Markdown文本分割成块
        返回 [{'type': block_type, 'content': content, 'start_line': start_line}, ...]
        """
        blocks = []
        lines = text.split('\n')
        
        i = 0
        while i < len(lines):
            line = lines[i]
            
            # 空行
            if not line.strip():
                i += 1
                continue
            
            # 代码块 ```
            if line.strip().startswith('```'):
                start_line = i + 1
                block, i = self._extract_code_block(lines, i)
                blocks.append({'type': 'code', 'content': block, 'start_line': start_line})
                continue
            
            # 表格
            # 支持表格前一行 {#tbl:...} 作为表格锚点
            if re.match(r'^\s*\{#([A-Za-z0-9_\-\:\.]+)\}\s*$', line.strip()) and i + 1 < len(lines) and '|' in lines[i + 1] and i + 2 < len(lines) and re.match(r'^[\s\|\:\-]+$', lines[i + 2]):
                start_line = i + 1
                m = re.match(r'^\s*\{#([A-Za-z0-9_\-\:\.]+)\}\s*$', line.strip())
                table_anchor = m.group(1) if m else None
                i += 1
                line = lines[i]
                block, i = self._extract_table(lines, i)
                blocks.append({'type': 'table', 'content': {'text': block, 'anchor': table_anchor}, 'start_line': start_line})
                continue

            if '|' in line and i + 1 < len(lines) and re.match(r'^[\s\|\:\-]+$', lines[i + 1]):
                start_line = i + 1
                block, i = self._extract_table(lines, i)
                blocks.append({'type': 'table', 'content': {'text': block, 'anchor': None}, 'start_line': start_line})
                continue
            
            # 块级公式 $$
            if line.strip().startswith('$$'):
                start_line = i + 1
                result = self._extract_block_math(lines, i)
                if len(result) == 3:
                    block, i, remaining = result
                    if remaining:
                        # 如果有剩余部分，插入到当前位置继续处理
                        lines = lines[:i] + [remaining] + lines[i:]
                else:
                    block, i = result
                blocks.append({'type': 'math_block', 'content': block, 'start_line': start_line})
                continue
            
            # 标题
            if line.startswith('#'):
                start_line = i + 1
                level = len(re.match(r'^#+', line).group())
                content = line[level:].strip()
                content, anchor = self._extract_anchor_suffix(content)
                blocks.append({'type': f'heading_{min(level, 6)}', 'content': {'text': content, 'anchor': anchor}, 'start_line': start_line})
                i += 1
                continue
            
            # 引用块
            if line.startswith('>'):
                start_line = i + 1
                block, i = self._extract_quote(lines, i)
                blocks.append({'type': 'quote', 'content': block, 'start_line': start_line})
                continue
            
            # 无序列表
            if re.match(r'^[\s]*[\*\-\+]\s', line):
                start_line = i + 1
                block, i = self._extract_list(lines, i, ordered=False)
                blocks.append({'type': 'bullet_list', 'content': block, 'start_line': start_line})
                continue
            
            # 有序列表
            if re.match(r'^[\s]*\d+\.\s', line):
                start_line = i + 1
                block, i = self._extract_list(lines, i, ordered=True)
                blocks.append({'type': 'numbered_list', 'content': block, 'start_line': start_line})
                continue
            
            # 水平分割线
            if re.match(r'^[\s]*[-\*_]{3,}[\s]*$', line):
                blocks.append({'type': 'hr', 'content': '', 'start_line': i + 1})
                i += 1
                continue
            
            # 图片（单独一行）
            img_match = re.match(r'^!\[([^\]]*)\]\(([^\)]+)\)\s*(\{#([A-Za-z0-9_\-\:\.]+)\})?\s*$', line.strip())
            if img_match:
                blocks.append({'type': 'image', 'content': {'text': (img_match.group(1), img_match.group(2)), 'anchor': (img_match.group(4) if img_match.group(4) else None)}, 'start_line': i + 1})
                i += 1
                continue

            # 目录标记 [[TOC]]（单独一行）
            if line.strip() == '[[TOC]]':
                blocks.append({'type': 'toc', 'content': '', 'start_line': i + 1})
                i += 1
                continue
            
            # 脚注定义 [^1]: 内容
            footnote_match = re.match(r'^\s*\[\^([^\]]+)\][:：]\s*(.*)', line.strip())
            if footnote_match:
                ref = footnote_match.group(1)
                content = footnote_match.group(2)
                blocks.append({'type': 'footnote_def', 'content': {'ref': ref, 'text': content}, 'start_line': i + 1})
                i += 1
                continue
            
            # 普通段落
            start_line = i + 1
            para, i = self._extract_paragraph(lines, i)
            blocks.append({'type': 'paragraph', 'content': para, 'start_line': start_line})
        
        return blocks

    def _extract_anchor_suffix(self, text: str) -> tuple[str, str]:
        """提取行尾的 {#id}。"""
        if not text:
            return text, None
        m = re.match(r'^(.*?)(?:\s*\{#([A-Za-z0-9_\-\:\.]+)\}\s*)$', text)
        if not m:
            return text, None
        body = (m.group(1) or '').rstrip()
        anchor = (m.group(2) or '').strip() or None
        return body, anchor

    def _sanitize_bookmark_name(self, anchor_id: str) -> str:
        s = str(anchor_id or '').strip()
        s = re.sub(r'[^A-Za-z0-9_]', '_', s)
        if not s:
            s = f'bkm_{self._bookmark_id_counter}'
        if not re.match(r'^[A-Za-z]', s):
            s = 'b_' + s
        return s[:40]

    def _ensure_bookmark(self, anchor_id: str) -> str:
        if not anchor_id:
            return None
        if anchor_id in self._id_to_bookmark:
            return self._id_to_bookmark[anchor_id]
        name = self._sanitize_bookmark_name(anchor_id)
        self._id_to_bookmark[anchor_id] = name
        return name

    def _add_bookmark_to_paragraph(self, paragraph, anchor_id: str) -> None:
        name = self._ensure_bookmark(anchor_id)
        if not name:
            return
        try:
            bid = str(self._bookmark_id_counter)
            self._bookmark_id_counter += 1

            start = OxmlElement('w:bookmarkStart')
            start.set(qn('w:id'), bid)
            start.set(qn('w:name'), name)
            end = OxmlElement('w:bookmarkEnd')
            end.set(qn('w:id'), bid)

            p = paragraph._p
            p.insert(0, start)
            p.append(end)
        except Exception:
            pass

    def _make_label_for_anchor(self, anchor_id: str, element_kind: str, heading_level: int = None) -> str:
        if not anchor_id:
            return None
        if anchor_id in self._id_to_label:
            return self._id_to_label[anchor_id]

        label = None
        try:
            if str(anchor_id).startswith('fig:'):
                self._fig_counter += 1
                label = f"图 {self._fig_counter}"
            elif str(anchor_id).startswith('tbl:'):
                self._tbl_counter += 1
                label = f"表 {self._tbl_counter}"
            elif str(anchor_id).startswith('sec:') and heading_level is not None:
                lv = max(1, min(4, int(heading_level)))
                self._sec_counters[lv - 1] += 1
                for j in range(lv, 4):
                    self._sec_counters[j] = 0
                parts = [str(self._sec_counters[k]) for k in range(lv) if self._sec_counters[k] > 0]
                num = '.'.join(parts) if parts else str(self._sec_counters[0])
                label = f"节 {num}"
        except Exception:
            label = None

        if not label:
            # 兜底：未知前缀
            label = str(anchor_id)

        self._id_to_label[anchor_id] = label
        return label

    def _add_internal_ref(self, paragraph, ref_id: str) -> None:
        rid = (ref_id or '').strip()
        if not rid:
            return

        bm = self._id_to_bookmark.get(rid)
        label = self._id_to_label.get(rid)
        if not label:
            # 未知引用：根据前缀给一个更友好的占位
            try:
                if rid.startswith('fig:'):
                    label = f"图 ?"
                elif rid.startswith('tbl:'):
                    label = f"表 ?"
                elif rid.startswith('sec:'):
                    label = f"节 ?"
                else:
                    label = rid
            except Exception:
                label = rid

        if not bm:
            # 未解析：红字占位 + 记录诊断
            try:
                run = paragraph.add_run(f"[未解析引用: {rid}]")
                run.font.color.rgb = RGBColor(255, 0, 0)
            except Exception:
                pass
            try:
                self.diagnostics['unresolved_refs'].append(rid)
            except Exception:
                pass
            return

        # 生成可点击的内部链接
        try:
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('w:anchor'), bm)
            hyperlink.set(qn('w:history'), '1')

            new_run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')

            color_val = str(self.export_style.get('hyperlink_color', '0000FF')).lstrip('#').upper()
            color = OxmlElement('w:color')
            color.set(qn('w:val'), color_val)
            rPr.append(color)

            underline = OxmlElement('w:u')
            underline_val = 'single' if bool(self.export_style.get('hyperlink_underline', True)) else 'none'
            underline.set(qn('w:val'), underline_val)
            rPr.append(underline)

            new_run.append(rPr)
            text_elem = OxmlElement('w:t')
            text_elem.text = str(label)
            new_run.append(text_elem)
            hyperlink.append(new_run)
            paragraph._p.append(hyperlink)
        except Exception:
            try:
                paragraph.add_run(str(label))
            except Exception:
                pass
    
    def _extract_code_block(self, lines: list, start: int) -> tuple:
        """提取代码块"""
        first_line = lines[start].strip()
        language = first_line[3:].strip()  # 获取语言标识
        
        code_lines = []
        i = start + 1
        
        while i < len(lines):
            if lines[i].strip() == '```':
                i += 1
                break
            code_lines.append(lines[i])
            i += 1
        
        return {'language': language, 'code': '\n'.join(code_lines)}, i
    
    def _extract_table(self, lines: list, start: int) -> tuple:
        """提取表格"""
        table_lines = []
        i = start
        
        while i < len(lines) and '|' in lines[i]:
            table_lines.append(lines[i])
            i += 1
        
        return '\n'.join(table_lines), i
    
    def _extract_block_math(self, lines: list, start: int) -> tuple:
        """提取块级公式"""
        first_line = lines[start].strip()
        
        # 单行公式 $$ ... $$ - 使用正则匹配第一个完整公式
        single_match = re.match(r'^\$\$(.+?)\$\$', first_line)
        if single_match:
            formula = single_match.group(1).strip()
            # 检查这行剩余部分是否还有公式
            remaining = first_line[single_match.end():].strip()
            if remaining.startswith('$$'):
                # 剩余部分有其他公式，创建副本而不是修改原列表
                return formula, start, remaining  # 返回剩余部分供调用者处理
            return formula, start + 1, None
        
        # 多行公式（以单独 $$ 开头）
        math_lines = []
        i = start + 1
        
        while i < len(lines):
            line_content = lines[i].strip()
            if line_content.endswith('$$'):
                if line_content != '$$':
                    math_lines.append(line_content[:-2])
                i += 1
                break
            math_lines.append(lines[i])
            i += 1
        
        return '\n'.join(math_lines).strip(), i, None
    
    def _extract_quote(self, lines: list, start: int) -> tuple:
        """提取引用块"""
        quote_lines = []
        i = start
        
        while i < len(lines) and lines[i].startswith('>'):
            # 移除 > 前缀
            content = lines[i][1:].strip() if len(lines[i]) > 1 else ''
            quote_lines.append(content)
            i += 1
        
        return '\n'.join(quote_lines), i
    
    def _extract_list(self, lines: list, start: int, ordered: bool) -> tuple:
        """提取列表（支持任务列表）"""
        items = []
        i = start
        
        pattern = r'^(\s*)(\d+\.|\*|\-|\+)\s+(.*)$'
        
        while i < len(lines):
            match = re.match(pattern, lines[i])
            if not match:
                break
            
            indent = len(match.group(1))
            level = indent // 2  # 每2个空格为一级
            text = match.group(3)
            
            # 检查是否是任务列表
            task_match = re.match(r'^\[([ xX])\]\s*(.*)', text)
            if task_match:
                checked = task_match.group(1).lower() == 'x'
                task_text = task_match.group(2)
                items.append((level, {'type': 'task', 'checked': checked, 'text': task_text}))
            else:
                items.append((level, {'type': 'item', 'text': text}))
            i += 1
        
        return items, i
    
    def _extract_paragraph(self, lines: list, start: int) -> tuple:
        """提取段落"""
        para_lines = []
        i = start
        
        while i < len(lines):
            line = lines[i]
            
            # 遇到块级元素则停止
            if not line.strip():
                break
            if line.startswith('#'):
                break
            if line.strip().startswith('```'):
                break
            if line.strip().startswith('$$'):
                break
            if line.startswith('>'):
                break
            if re.match(r'^[\s]*[\*\-\+]\s', line):
                break
            if re.match(r'^[\s]*\d+\.\s', line):
                break
            if re.match(r'^[\s]*[-\*_]{3,}[\s]*$', line):
                break
            if '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
                # 可能是表格
                if re.match(r'^[\s\|\:\-]+$', lines[i + 1]):
                    break
            
            para_lines.append(line)
            i += 1
        
        return ' '.join(para_lines), i
    
    def _setup_footnotes_part(self):
        """初始化脚注系统 - 使用延迟 ZIP 注入方式。
        
        由于 python-docx 不支持动态添加 footnotes.xml 部件，
        我们将脚注信息收集起来，在 save() 时通过 ZIP 操作注入。
        """
        # 如果已初始化则跳过
        if hasattr(self, '_footnotes_initialized') and self._footnotes_initialized:
            return
        
        self._footnote_id_counter = 2  # 从 2 开始（0 和 1 被保留）
        self._footnote_ref_to_id = {}  # 映射 Markdown 引用 ID 到 Word 脚注 ID
        self._pending_footnotes = []  # 待注入的脚注内容列表
        self._footnotes_initialized = True
    
    def _add_native_footnote(self, paragraph, ref_id: str, content: str):
        """向段落中插入原生 Word 脚注引用。
        
        实际的脚注内容会在文档保存时通过 ZIP 注入。
        """
        # 确保脚注系统已初始化
        self._setup_footnotes_part()
        
        # 分配一个 Word 脚注 ID
        if ref_id not in self._footnote_ref_to_id:
            fn_id = str(self._footnote_id_counter)
            self._footnote_id_counter += 1
            self._footnote_ref_to_id[ref_id] = fn_id
            
            # 记录待注入的脚注
            self._pending_footnotes.append({
                'id': fn_id,
                'content': content or ''
            })
        else:
            fn_id = self._footnote_ref_to_id[ref_id]
        
        # 在正文段落中插入脚注引用
        run = paragraph.add_run()
        run._r.append(self._create_footnote_reference_element(fn_id))
    
    def _create_footnote_reference_element(self, fn_id: str):
        """创建 w:footnoteReference 元素"""
        footnote_ref = OxmlElement('w:footnoteReference')
        footnote_ref.set(qn('w:id'), fn_id)
        return footnote_ref
    
    def _inject_footnotes_into_docx(self, docx_path: str):
        """在保存后的 docx 文件中注入 footnotes.xml。
        
        这是一个 ZIP 操作，直接修改 docx 文件内容。
        """
        import zipfile
        import tempfile
        import shutil
        from lxml import etree
        
        if not hasattr(self, '_pending_footnotes') or not self._pending_footnotes:
            return
        
        # 创建 footnotes.xml 内容
        nsmap = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        }
        w = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        
        footnotes_root = etree.Element(f'{w}footnotes', nsmap=nsmap)
        
        # 分隔符脚注 (id=0)
        fn_sep = etree.SubElement(footnotes_root, f'{w}footnote', {f'{w}type': 'separator', f'{w}id': '0'})
        p_sep = etree.SubElement(fn_sep, f'{w}p')
        r_sep = etree.SubElement(p_sep, f'{w}r')
        etree.SubElement(r_sep, f'{w}separator')
        
        # 续页分隔符脚注 (id=1)
        fn_cont = etree.SubElement(footnotes_root, f'{w}footnote', {f'{w}type': 'continuationSeparator', f'{w}id': '1'})
        p_cont = etree.SubElement(fn_cont, f'{w}p')
        r_cont = etree.SubElement(p_cont, f'{w}r')
        etree.SubElement(r_cont, f'{w}continuationSeparator')
        
        # 添加实际脚注
        for fn in self._pending_footnotes:
            fn_elem = etree.SubElement(footnotes_root, f'{w}footnote', {f'{w}id': fn['id']})
            p_elem = etree.SubElement(fn_elem, f'{w}p')
            
            # 脚注引用标记
            r_ref = etree.SubElement(p_elem, f'{w}r')
            rPr_ref = etree.SubElement(r_ref, f'{w}rPr')
            etree.SubElement(rPr_ref, f'{w}rStyle', {f'{w}val': 'FootnoteReference'})
            etree.SubElement(r_ref, f'{w}footnoteRef')
            
            # 空格
            r_space = etree.SubElement(p_elem, f'{w}r')
            t_space = etree.SubElement(r_space, f'{w}t')
            t_space.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            t_space.text = ' '
            
            # 脚注内容
            r_text = etree.SubElement(p_elem, f'{w}r')
            t_text = etree.SubElement(r_text, f'{w}t')
            t_text.text = fn['content']
        
        footnotes_xml = etree.tostring(footnotes_root, xml_declaration=True, encoding='UTF-8', standalone='yes')
        
        # 读取并修改 docx (ZIP)
        temp_path = docx_path + '.tmp'
        
        with zipfile.ZipFile(docx_path, 'r') as zin:
            with zipfile.ZipFile(temp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    data = zin.read(item.filename)
                    
                    # 更新 [Content_Types].xml
                    if item.filename == '[Content_Types].xml':
                        data = self._update_content_types(data)
                    
                    # 更新 word/_rels/document.xml.rels
                    elif item.filename == 'word/_rels/document.xml.rels':
                        data = self._update_document_rels(data)
                    
                    zout.writestr(item, data)
                
                # 添加 footnotes.xml
                zout.writestr('word/footnotes.xml', footnotes_xml)
        
        # 替换原文件
        shutil.move(temp_path, docx_path)
    
    def _update_content_types(self, data: bytes) -> bytes:
        """更新 [Content_Types].xml 以包含 footnotes.xml"""
        from lxml import etree
        
        root = etree.fromstring(data)
        ns = '{http://schemas.openxmlformats.org/package/2006/content-types}'
        
        # 检查是否已存在
        for override in root.findall(f'{ns}Override'):
            if override.get('PartName') == '/word/footnotes.xml':
                return data
        
        # 添加 Override
        override = etree.SubElement(root, f'{ns}Override')
        override.set('PartName', '/word/footnotes.xml')
        override.set('ContentType', 'application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml')
        
        return etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone='yes')
    
    def _update_document_rels(self, data: bytes) -> bytes:
        """更新 word/_rels/document.xml.rels 以引用 footnotes.xml"""
        from lxml import etree
        
        root = etree.fromstring(data)
        ns = '{http://schemas.openxmlformats.org/package/2006/relationships}'
        
        # 检查是否已存在
        for rel in root.findall(f'{ns}Relationship'):
            if rel.get('Target') == 'footnotes.xml':
                return data
        
        # 找到最大 rId
        max_id = 0
        for rel in root.findall(f'{ns}Relationship'):
            rid = rel.get('Id', '')
            if rid.startswith('rId'):
                try:
                    num = int(rid[3:])
                    max_id = max(max_id, num)
                except ValueError:
                    pass
        
        # 添加 Relationship
        rel = etree.SubElement(root, f'{ns}Relationship')
        rel.set('Id', f'rId{max_id + 1}')
        rel.set('Type', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes')
        rel.set('Target', 'footnotes.xml')
        
        return etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone='yes')



    def _apply_header_footer(self):
        """应用页眉页脚配置"""
        config = self.header_footer_config
        if not config:
            return
            
        section = self.doc.sections[0]
        
        # 页眉
        if config.get('header_text') or config.get('show_date'):
            header = section.header
            # 清除已有内容
            for paragraph in header.paragraphs:
                p = paragraph._element
                p.getparent().remove(p)
                
            p = header.add_paragraph()
            text_parts = []
            if config.get('header_text'):
                text_parts.append(config.get('header_text'))
            if config.get('show_date'):
                import datetime
                text_parts.append(datetime.datetime.now().strftime("%Y-%m-%d"))
                
            p.text = " | ".join(text_parts)
            
            align = config.get('header_align', 'center')
            if align == 'left':
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif align == 'right':
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
        # 页脚
        if config.get('footer_text') or config.get('show_page_num'):
            footer = section.footer
            # 清除已有内容
            for paragraph in footer.paragraphs:
                p = paragraph._element
                p.getparent().remove(p)
                
            p = footer.add_paragraph()
            
            # 页码处理比较复杂，这里简化处理，只添加文本
            text_parts = []
            if config.get('footer_text'):
                text_parts.append(config.get('footer_text'))
                
            p.text = " ".join(text_parts)
            
            # 如果需要显示页码
            if config.get('show_page_num'):
                if p.text:
                    run = p.add_run("  ")
                self._add_page_number(p)
            
            align = config.get('footer_align', 'center')
            if align == 'left':
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif align == 'right':
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_page_number(self, paragraph):
        """向段落添加页码字段"""
        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
    
    def _process_block(self, block_type: str, content) -> None:
        """处理单个块"""
        
        if block_type.startswith('heading_'):
            level = int(block_type.split('_')[1])
            anchor = content.get('anchor') if isinstance(content, dict) else None
            text = content.get('text') if isinstance(content, dict) else content
            self._add_heading(text, level, anchor_id=anchor)
        
        elif block_type == 'paragraph':
            self._add_paragraph(content)
        
        elif block_type == 'image':
            anchor = content.get('anchor') if isinstance(content, dict) else None
            alt_text, image_path = (content.get('text') if isinstance(content, dict) else content)
            ok = self.image_handler.add_image(self.doc, image_path, alt_text)

            # 给图片添加书签（用于引用）
            if anchor:
                try:
                    # 图片插入时会创建一个居中段落，取最后一个段落作为锚点
                    p = self.doc.paragraphs[-1] if self.doc.paragraphs else None
                    if p is not None:
                        self._add_bookmark_to_paragraph(p, anchor)
                        self._make_label_for_anchor(anchor, 'image')
                except Exception:
                    pass
            
        elif block_type == 'footnote_def':
            # 脚注定义，暂存不直接渲染
            # content 现在是一个字典 {'ref': ..., 'text': ...}
            if isinstance(content, dict):
                ref = content.get('ref')
                text = content.get('text')
                if ref:
                    self.footnotes[ref] = text
        
        elif block_type == 'code':
            self.code_handler.add_code_block(
                self.doc, 
                content['code'], 
                content.get('language')
            )
        
        elif block_type == 'table':
            table_anchor = content.get('anchor') if isinstance(content, dict) else None
            table_text = content.get('text') if isinstance(content, dict) else content

            headers, rows, alignments = self.table_handler.parse_markdown_table(table_text)
            if headers:
                # 使用自定义方法处理表格，支持单元格内公式
                self._add_table_with_formulas(headers, rows, alignments, anchor_id=table_anchor)
        
        elif block_type == 'math_block':
            self.math_handler.add_block_equation(self.doc, content)
        
        elif block_type == 'quote':
            para = self.doc.add_paragraph()
            if self._use_word_styles():
                self._apply_paragraph_style(para, self._mapped_style_name('map_quote', 'Block Quote'))
            else:
                try:
                    para.style = 'Block Quote'
                except Exception:
                    pass
            # 使用行内元素解析
            self._add_inline_content(para, content)
            for run in para.runs:
                self._set_body_font(run, keep_style=True)
        
        elif block_type == 'bullet_list':
            for level, item in content:
                # 使用行内元素解析
                self._add_list_item(item, level, ordered=False)
        
        elif block_type == 'numbered_list':
            # 为每个新的有序列表重置编号
            for idx, (level, item) in enumerate(content):
                # 第一项需要重新开始编号
                self._add_list_item(item, level, ordered=True, restart=(idx == 0))
        
        elif block_type == 'image':
            anchor = content.get('anchor') if isinstance(content, dict) else None
            alt_text, image_path = (content.get('text') if isinstance(content, dict) else content)
            ok = self.image_handler.add_image(self.doc, image_path, alt_text)

            # 给图片添加书签（用于引用）
            if anchor:
                try:
                    # 图片插入时会创建一个居中段落，取最后一个段落作为锚点
                    p = self.doc.paragraphs[-1] if self.doc.paragraphs else None
                    if p is not None:
                        self._add_bookmark_to_paragraph(p, anchor)
                        self._make_label_for_anchor(anchor, 'image')
                except Exception:
                    pass

        elif block_type == 'toc':
            if bool(self.export_style.get('toc_enabled', False)):
                self._add_toc()
        
        elif block_type == 'hr':
            # 添加水平分割线（使用段落底部边框）
            para = self.doc.add_paragraph()
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            
            # 使用段落底部边框作为分割线
            p = para._p
            pPr = p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')  # 线宽（半磅）
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '808080')  # 灰色
            pBdr.append(bottom)
            pPr.append(pBdr)

    def _add_toc(self) -> None:
        """在当前位置插入 Word 目录域（TOC）。"""
        try:
            para = self.doc.add_paragraph()
            run = para.add_run()

            fld_begin = OxmlElement('w:fldChar')
            fld_begin.set(qn('w:fldCharType'), 'begin')

            instr = OxmlElement('w:instrText')
            instr.set(qn('xml:space'), 'preserve')
            # 1-3 级标题；支持超链接；隐藏页码在 web 视图的提示等
            instr.text = 'TOC \\o "1-3" \\h \\z \\u'

            fld_sep = OxmlElement('w:fldChar')
            fld_sep.set(qn('w:fldCharType'), 'separate')

            # placeholder
            t = OxmlElement('w:t')
            t.text = '目录将在打开文档后自动生成/更新'

            fld_end = OxmlElement('w:fldChar')
            fld_end.set(qn('w:fldCharType'), 'end')

            r = run._r
            r.append(fld_begin)
            r.append(instr)
            r.append(fld_sep)
            r.append(t)
            r.append(fld_end)
        except Exception:
            # 失败时不抛异常，避免影响导出
            pass
    
    def _add_heading(self, text: str, level: int, anchor_id: str = None) -> None:
        """添加标题 - 支持公式等行内元素解析"""
        # 标题配置：(字号, 居中)
        heading_config = {
            1: (FONT_SIZES['er_hao'], True),    # 一级：二号，居中
            2: (FONT_SIZES['san_hao'], True),   # 二级：三号，居中
            3: (FONT_SIZES['xiao_san'], False), # 三级：小三，靠左
            4: (FONT_SIZES['si_hao'], False),   # 四级：四号，靠左
        }
        
        level = min(level, 4)
        font_size, is_center = heading_config.get(level, heading_config[4])
        
        # 创建段落
        para = self.doc.add_paragraph()

        if self._use_word_styles():
            style_name = self._mapped_style_name(f'map_heading_{level}', f'Heading {level}')
            self._apply_paragraph_style(para, style_name)
        
        if not self._use_word_styles():
            # 设置对齐方式
            if is_center:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # 设置段前段后间距
            para.paragraph_format.space_before = Pt(18 if level == 1 else 12)
            para.paragraph_format.space_after = Pt(12 if level == 1 else 8)
        
        # 使用行内元素解析（支持公式等）
        self._add_inline_content(para, text)

        # 方案B：使用 Word 内置标题样式时，某些 run（加粗/斜体/超链接等直接格式）
        # 可能会丢失 eastAsia 字体继承，导致中文标题显示为默认字体。
        # 这里为标题段落的每个 run 补齐中文字体，避免标题字体“跳变”。
        if self._use_word_styles():
            try:
                self._ensure_heading_paragraph_fonts(para)
            except Exception:
                pass

        # 书签（用于引用）+ 节编号标签
        if anchor_id:
            try:
                self._add_bookmark_to_paragraph(para, anchor_id)
                if str(anchor_id).startswith('sec:'):
                    self._make_label_for_anchor(anchor_id, 'heading', heading_level=level)
            except Exception:
                pass
        
        if not self._use_word_styles():
            # 为所有run设置标题字体样式
            for run in para.runs:
                self._set_heading_font(run, font_size)
    
    def _ensure_heading_paragraph_fonts(self, paragraph) -> None:
        """方案B：确保标题段落内所有 run（含超链接等）都补齐中文字体。

        - Word 中超链接通常是 w:hyperlink/w:r，不一定出现在 paragraph.runs 中。
        - 标题含引号等字符时，Word/库可能拆分为多个 run。
        - 这里只补齐 eastAsia，避免覆盖字号/加粗等由样式控制的属性。
        """
        try:
            heading_cn = self.export_style.get('heading_cn', FONTS['heading_cn'])
        except Exception:
            heading_cn = FONTS['heading_cn']

        try:
            p = paragraph._p
            # 遍历段落下所有 w:r（包含 hyperlink 内部的 run）
            runs = p.xpath('.//w:r')
        except Exception:
            runs = []

        for r in runs:
            try:
                rPr = r.get_or_add_rPr()
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = OxmlElement('w:rFonts')
                    rPr.insert(0, rFonts)
                rFonts.set(qn('w:eastAsia'), heading_cn)
            except Exception:
                pass

    def _set_heading_font(self, run, font_size) -> None:
        """设置标题字体 - 中文黑体，英文Times New Roman，加粗"""
        if self._use_word_styles():
            return
        run.font.size = font_size
        run.font.bold = True
        run.font.name = FONTS['body_en']  # 英文字体
        
        # 设置中文字体为黑体
        r = run._r
        rPr = r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), FONTS['heading_cn'])  # 黑体
    
    def _add_paragraph(self, text: str) -> None:
        """添加段落 - 正文：中文宋体小四，英文Times New Roman小四，首行缩进2字符"""
        if not text.strip():
            return
        
        para = self.doc.add_paragraph()
        if self._use_word_styles():
            self._apply_paragraph_style(para, self._mapped_style_name('map_paragraph', 'Normal'))
        else:
            # 设置首行缩进：2个字符（小四12pt × 2 = 24pt）
            para.paragraph_format.first_line_indent = Pt(24)
        
        self._add_inline_content(para, text)
        
        # 为段落中的所有run设置正文字体（保留原有的bold/italic状态）
        for run in para.runs:
            self._set_body_font(run, keep_style=True)
    
    def _add_list_item(self, item, level: int = 0, ordered: bool = False, restart: bool = False) -> None:
        """添加列表项 - 支持任务列表、行内元素解析和编号重置
        
        Args:
            item: 列表项内容
            level: 缩进级别（0=顶级, 1=子级, 2=孙级...）
            ordered: 是否为有序列表
            restart: 是否重新开始编号（列表第一项）
        """
        # 处理新格式（字典）和旧格式（字符串）
        if isinstance(item, dict):
            item_type = item.get('type', 'item')
            text = item.get('text', '')
            is_task = item_type == 'task'
            checked = item.get('checked', False)
        else:
            text = item
            is_task = False
            checked = False
        
        if is_task:
            # 任务列表项：使用普通段落 + checkbox符号
            para = self.doc.add_paragraph()
            checkbox = '☑' if checked else '☐'
            run = para.add_run(f"{checkbox} ")
            run.font.size = FONT_SIZES['xiao_si']
            self._add_inline_content(para, text)
        else:
            # 普通列表项
            if ordered:
                # 有序列表：手动管理编号
                para = self.doc.add_paragraph()
                self._apply_list_numbering(para, level, restart)
            else:
                para = self.doc.add_paragraph(style='List Bullet')
            self._add_inline_content(para, text)
        
        # 设置字体
        for run in para.runs:
            self._set_body_font(run, keep_style=True)
    
    def _apply_list_numbering(self, paragraph, level: int = 0, restart: bool = False) -> None:
        """为段落应用有序列表编号
        
        Args:
            paragraph: 段落对象
            level: 列表级别（0=顶级, 1=子级...)
            restart: 是否重新开始编号（新列表的第一项）
        """
        if restart:
            # 为新列表创建新的编号定义
            self._current_num_id = self._create_numbering_definition()
        
        # 应用编号到段落
        p = paragraph._p
        pPr = p.get_or_add_pPr()
        
        numPr = OxmlElement('w:numPr')
        ilvl = OxmlElement('w:ilvl')
        ilvl.set(qn('w:val'), str(level))  # 使用实际级别
        numPr.append(ilvl)
        
        numId = OxmlElement('w:numId')
        numId.set(qn('w:val'), str(self._current_num_id))
        numPr.append(numId)
        
        pPr.insert(0, numPr)
    
    def _create_numbering_definition(self) -> int:
        """创建新的多级编号定义，返回 numId
        
        级别0: 1. 2. 3. (数字)
        级别1: a) b) c) (小写字母)
        级别2: i. ii. iii. (小写罗马数字)
        """
        num_id = self._num_id_counter
        self._num_id_counter += 1
        abstract_num_id = num_id + 100
        
        # 多级列表配置: (格式, 文本模板, 缩进左, 悬挂缩进)
        level_configs = [
            ('decimal', '%1.', 720, 360),      # 1. 2. 3.
            ('lowerLetter', '%2)', 1080, 360), # a) b) c)
            ('lowerRoman', '%3.', 1440, 360),  # i. ii. iii.
            ('decimal', '%4.', 1800, 360),     # 1. 2. 3.
            ('lowerLetter', '%5)', 2160, 360), # a) b) c)
        ]
        
        try:
            numbering_part = self.doc.part.numbering_part
            numbering = numbering_part.numbering_definitions._numbering
            
            # 创建 abstractNum
            abstract_num = OxmlElement('w:abstractNum')
            abstract_num.set(qn('w:abstractNumId'), str(abstract_num_id))
            
            # 多级别支持标记
            multi_level = OxmlElement('w:multiLevelType')
            multi_level.set(qn('w:val'), 'multilevel')
            abstract_num.append(multi_level)
            
            # 创建多个级别定义
            for i, (num_fmt, lvl_text, left_indent, hanging) in enumerate(level_configs):
                lvl = OxmlElement('w:lvl')
                lvl.set(qn('w:ilvl'), str(i))
                
                start = OxmlElement('w:start')
                start.set(qn('w:val'), '1')
                lvl.append(start)
                
                numFmt = OxmlElement('w:numFmt')
                numFmt.set(qn('w:val'), num_fmt)
                lvl.append(numFmt)
                
                lvlText = OxmlElement('w:lvlText')
                lvlText.set(qn('w:val'), lvl_text)
                lvl.append(lvlText)
                
                lvlJc = OxmlElement('w:lvlJc')
                lvlJc.set(qn('w:val'), 'left')
                lvl.append(lvlJc)
                
                # 缩进设置
                pPr = OxmlElement('w:pPr')
                ind = OxmlElement('w:ind')
                ind.set(qn('w:left'), str(left_indent))
                ind.set(qn('w:hanging'), str(hanging))
                pPr.append(ind)
                lvl.append(pPr)
                
                abstract_num.append(lvl)
            
            # 插入 abstractNum
            numbering.insert(0, abstract_num)
            
            # 创建 num 元素
            num = OxmlElement('w:num')
            num.set(qn('w:numId'), str(num_id))
            abstractNumId_ref = OxmlElement('w:abstractNumId')
            abstractNumId_ref.set(qn('w:val'), str(abstract_num_id))
            num.append(abstractNumId_ref)
            numbering.append(num)
            
        except Exception as e:
            # 如果失败，使用默认值
            pass
        
        return num_id
    
    def _add_table_with_formulas(self, headers: list, rows: list, alignments: list = None, anchor_id: str = None) -> None:
        """添加表格 - 支持单元格内公式渲染，支持对齐方式，支持 {#tbl:...} 书签"""
        from styles import apply_table_style
        
        num_cols = len(headers)
        num_rows = len(rows) + 1
        
        # 表格锚点：在表格前插入一个空段落作为书签位置
        if anchor_id:
            try:
                p = self.doc.add_paragraph()
                self._add_bookmark_to_paragraph(p, anchor_id)
                if str(anchor_id).startswith('tbl:'):
                    self._make_label_for_anchor(anchor_id, 'table')
            except Exception:
                pass

        # 创建表格
        table = self.doc.add_table(rows=num_rows, cols=num_cols)

        # 表格宽度：按页面可用宽度等分列宽，避免超出页边距
        try:
            if self.doc.sections:
                sec = self.doc.sections[0]
                available = sec.page_width - sec.left_margin - sec.right_margin
                if num_cols > 0 and available is not None:
                    table.autofit = False
                    col_w = int(available / num_cols)
                    for i in range(num_cols):
                        try:
                            table.columns[i].width = col_w
                        except Exception:
                            pass
                        try:
                            for c in table.columns[i].cells:
                                c.width = col_w
                        except Exception:
                            pass
        except Exception:
            pass
        
        # 填充表头（表头始终居中）
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            para = cell.paragraphs[0]
            para.clear()
            self._add_inline_content(para, header)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 表头始终居中
            # 设置字体
            for run in para.runs:
                self._set_body_font(run, keep_style=True)
                run.bold = True  # 表头加粗

        # A方案：所有表格默认第一行重复为表头（跨页自动重复）
        try:
            tr = table.rows[0]._tr
            trPr = tr.get_or_add_trPr()
            tblHeader = OxmlElement('w:tblHeader')
            tblHeader.set(qn('w:val'), 'true')
            trPr.append(tblHeader)
        except Exception:
            pass
        
        # 填充数据行（根据对齐方式设置）
        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < num_cols:
                    cell = table.cell(row_idx + 1, col_idx)
                    para = cell.paragraphs[0]
                    para.clear()
                    self._add_inline_content(para, cell_text)
                    # 根据对齐方式设置
                    if alignments and col_idx < len(alignments):
                        align = alignments[col_idx]
                        if align == 'center':
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif align == 'right':
                            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        else:
                            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    else:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 默认居中
                    # 设置字体
                    for run in para.runs:
                        self._set_body_font(run, keep_style=True)
        
        # 应用表格样式
        apply_table_style(table, self.export_style)
    
    def _add_inline_content(self, paragraph, text: str, base_bold: bool = False, base_italic: bool = False) -> None:
        """处理行内元素：使用共用解析器，支持嵌套公式解析
        
        Args:
            paragraph: 段落对象
            text: 要解析的文本
            base_bold: 基础加粗样式（用于嵌套调用）
            base_italic: 基础斜体样式（用于嵌套调用）
        """
        elements = parse_inline(text)
        
        ref_pat = re.compile(r'\[\[ref:([^\]]+)\]\]')

        for elem in elements:
            if elem.type == InlineType.TEXT:
                s = elem.content or ''
                last = 0
                for m in ref_pat.finditer(s):
                    if m.start() > last:
                        run = paragraph.add_run(s[last:m.start()])
                        if base_bold:
                            run.bold = True
                        if base_italic:
                            run.italic = True

                    ref_id = (m.group(1) or '').strip()
                    self._add_internal_ref(paragraph, ref_id)
                    last = m.end()

                if last < len(s):
                    run = paragraph.add_run(s[last:])
                    if base_bold:
                        run.bold = True
                    if base_italic:
                        run.italic = True
            
            elif elem.type == InlineType.BOLD:
                # 递归解析加粗内容（支持内部公式）
                self._add_inline_content(paragraph, elem.content, base_bold=True, base_italic=base_italic)
            
            elif elem.type == InlineType.ITALIC:
                # 递归解析斜体内容（支持内部公式）
                self._add_inline_content(paragraph, elem.content, base_bold=base_bold, base_italic=True)
            
            elif elem.type == InlineType.BOLD_ITALIC:
                # 递归解析粗斜体内容（支持内部公式）
                self._add_inline_content(paragraph, elem.content, base_bold=True, base_italic=True)
            
            elif elem.type == InlineType.CODE:
                self.code_handler.add_inline_code(paragraph, elem.content)
            
            elif elem.type == InlineType.MATH:
                if is_valid_math_formula(elem.content):
                    self.math_handler.add_inline_equation(paragraph, elem.content)
                else:
                    run = paragraph.add_run(f'${elem.content}$')
                    if base_bold:
                        run.bold = True
                    if base_italic:
                        run.italic = True
            
            elif elem.type == InlineType.LINK:
                # 添加可点击的超链接
                self._add_hyperlink(paragraph, elem.url, elem.content)
            
            elif elem.type == InlineType.IMAGE:
                if elem.url:
                    self.image_handler.add_inline_image(paragraph, elem.url, elem.content)
                else:
                    paragraph.add_run(f'[图片: {elem.content}]')
            
            elif elem.type == InlineType.STRIKETHROUGH:
                run = paragraph.add_run(elem.content)
                run.font.strike = True
                if base_bold:
                    run.bold = True
                if base_italic:
                    run.italic = True
            
            elif elem.type == InlineType.SUPERSCRIPT:
                # 上标
                run = paragraph.add_run(elem.content or '')
                run.font.superscript = True
                if base_bold: run.bold = True
                if base_italic: run.italic = True
                self._set_body_font(run, keep_style=True)
                
            elif elem.type == InlineType.SUBSCRIPT:
                # 下标
                run = paragraph.add_run(elem.content or '')
                run.font.subscript = True
                if base_bold: run.bold = True
                if base_italic: run.italic = True
                self._set_body_font(run, keep_style=True)

            elif elem.type == InlineType.FOOTNOTE_REF:
                # 脚注引用：使用原生 Word 脚注
                ref_id = elem.content
                # 查找对应的脚注内容
                fn_content = self.footnotes.get(ref_id, '')
                if fn_content or ref_id in self.footnotes:
                    self._add_native_footnote(paragraph, ref_id, fn_content)
                else:
                    # 如果定义尚未出现，先占位，稍后在 convert_text 中填充
                    # 暂时插入普通文本标记，之后处理
                    run = paragraph.add_run(f"[^{ref_id}]")
                    run.font.superscript = True
                    run.font.size = Pt(9)

            elif elem.type == InlineType.LINEBREAK:
                # 换行
                paragraph.add_run('\n')
    
    def _set_body_font(self, run, keep_style: bool = False) -> None:
        """设置正文字体 - 中文宋体小四，英文Times New Roman小四
        
        Args:
            run: 文本run对象
            keep_style: 是否保留原有的bold/italic样式
        """
        if self._use_word_styles():
            # 方案B：由 Word 样式决定字体/段落格式，这里只保留加粗/斜体等行内样式
            try:
                if keep_style:
                    if run.bold:
                        run.font.bold = True
                    if run.italic:
                        run.font.italic = True
                return
            except Exception:
                return

        # 先保存原有样式
        original_bold = run.bold
        original_italic = run.italic
        
        body_size = self.export_style.get('body_size_pt')
        run.font.size = Pt(float(body_size)) if body_size is not None else FONT_SIZES['xiao_si']
        run.font.name = self.export_style.get('body_en', FONTS['body_en'])
        
        # 根据参数决定是否保留原有样式
        if keep_style:
            # 保留原有的bold/italic
            if original_bold:
                run.font.bold = True
            if original_italic:
                run.font.italic = True
        else:
            run.font.bold = False
        
        # 设置中文字体为宋体
        r = run._r
        rPr = r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), self.export_style.get('body_cn', FONTS['body_cn']))
    
    def _add_hyperlink(self, paragraph, url: str, text: str) -> None:
        """添加可点击的超链接"""
        # 创建关系
        part = paragraph.part
        r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
        
        # 创建超链接元素
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        
        # 创建 run 元素
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        
        # 设置颜色和下划线
        color_val = str(self.export_style.get('hyperlink_color', '0000FF')).lstrip('#').upper()
        color = OxmlElement('w:color')
        color.set(qn('w:val'), color_val)
        rPr.append(color)
        
        underline = OxmlElement('w:u')
        underline_val = 'single' if bool(self.export_style.get('hyperlink_underline', True)) else 'none'
        underline.set(qn('w:val'), underline_val)
        rPr.append(underline)
        
        # 设置字体大小
        link_size_pt = self.export_style.get('hyperlink_size_pt', self.export_style.get('body_size_pt', 12))
        try:
            half_points = str(int(float(link_size_pt) * 2))
        except Exception:
            half_points = '24'
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), half_points)
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), half_points)
        rPr.append(szCs)
        
        new_run.append(rPr)
        
        # 添加文本
        text_elem = OxmlElement('w:t')
        text_elem.text = text
        new_run.append(text_elem)
        
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
    
    def save(self, output_path: str) -> None:
        """保存文档"""
        if self.doc:
            self.doc.save(output_path)
            # 注入脚注（如果有）
            try:
                self._inject_footnotes_into_docx(output_path)
            except Exception as e:
                print(f"脚注注入失败: {e}")
            print(f"文档已保存: {output_path}")

