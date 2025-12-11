# -*- coding: utf-8 -*-
"""
元素处理器模块 - 表格、公式、图片、代码块处理
"""

import re
import warnings
warnings.filterwarnings('ignore')

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
from latex2mathml.converter import convert as latex_to_mathml

from utils import resolve_image_path, get_image_dimensions, extract_latex_from_text
from styles import apply_table_style, FONTS, FONT_SIZES


class ImageHandler:
    """图片处理器"""
    
    def __init__(self, base_dir: str = None):
        self.base_dir = base_dir
        self.image_counter = 0
    
    def add_image(self, doc: Document, image_path: str, alt_text: str = "", 
                  caption: str = None, max_width: float = 6.0) -> bool:
        """
        添加图片到文档
        返回是否成功
        """
        resolved_path = resolve_image_path(image_path, self.base_dir)
        if not resolved_path:
            # 图片不存在，添加占位文本
            para = doc.add_paragraph()
            run = para.add_run(f"[图片无法加载: {image_path}]")
            run.font.color.rgb = RGBColor(255, 0, 0)
            return False
        
        try:
            # 获取适当的图片尺寸
            width_inches, height_inches = get_image_dimensions(resolved_path, max_width)
            
            # 创建居中段落
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加图片
            run = para.add_run()
            run.add_picture(resolved_path, width=Inches(width_inches))
            
            # 添加标题（如果有）
            if caption or alt_text:
                self.image_counter += 1
                caption_text = caption or alt_text
                cap_para = doc.add_paragraph()
                cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_run = cap_para.add_run(f"图 {self.image_counter}: {caption_text}")
                cap_run.font.size = FONT_SIZES['caption']
                cap_run.font.name = FONTS['body_en']
            
            return True
            
        except (IOError, OSError) as e:
            # 静默处理图片加载失败
            para = doc.add_paragraph()
            run = para.add_run(f"[图片加载错误: {image_path}]")
            run.font.color.rgb = RGBColor(255, 0, 0)
            return False


class TableHandler:
    """表格处理器"""
    
    def __init__(self):
        self.table_counter = 0
    
    def parse_markdown_table(self, table_text: str) -> tuple:
        """
        解析Markdown表格
        返回 (headers, rows, alignments)
        """
        lines = [l.strip() for l in table_text.strip().split('\n') if l.strip()]
        if len(lines) < 2:
            return None, None, None
        
        # 解析表头
        headers = self._parse_row(lines[0])
        
        # 解析对齐方式
        alignments = self._parse_alignments(lines[1])
        
        # 解析数据行
        rows = []
        for line in lines[2:]:
            row = self._parse_row(line)
            if row:
                rows.append(row)
        
        return headers, rows, alignments
    
    def _parse_row(self, line: str) -> list:
        """解析表格行"""
        # 移除首尾的 |
        line = line.strip()
        if line.startswith('|'):
            line = line[1:]
        if line.endswith('|'):
            line = line[:-1]
        
        cells = [cell.strip() for cell in line.split('|')]
        return cells
    
    def _parse_alignments(self, line: str) -> list:
        """解析对齐方式行"""
        cells = self._parse_row(line)
        alignments = []
        
        for cell in cells:
            cell = cell.strip()
            if cell.startswith(':') and cell.endswith(':'):
                alignments.append('center')
            elif cell.endswith(':'):
                alignments.append('right')
            elif cell.startswith(':'):
                alignments.append('left')
            else:
                alignments.append('center')  # 默认居中
        
        return alignments
    
    def add_table(self, doc: Document, headers: list, rows: list, 
                  alignments: list = None, caption: str = None) -> None:
        """添加表格到文档"""
        if not headers:
            return
        
        num_cols = len(headers)
        num_rows = len(rows) + 1  # +1 for header
        
        # 创建表格
        table = doc.add_table(rows=num_rows, cols=num_cols)
        
        # 填充表头
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            self._apply_cell_alignment(cell, alignments[i] if alignments and i < len(alignments) else 'left')
        
        # 填充数据行
        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < num_cols:
                    cell = table.cell(row_idx + 1, col_idx)
                    cell.text = cell_text
                    self._apply_cell_alignment(cell, alignments[col_idx] if alignments and col_idx < len(alignments) else 'left')
        
        # 应用LaTeX三线表样式
        apply_table_style(table)
        
        # 添加表格标题
        if caption:
            self.table_counter += 1
            cap_para = doc.add_paragraph()
            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_run = cap_para.add_run(f"表 {self.table_counter}: {caption}")
            cap_run.font.size = FONT_SIZES['caption']


class MathHandler:
    """公式处理器 - LaTeX公式转Word公式"""
    
    def __init__(self):
        self.equation_counter = 0
    
    def add_block_equation(self, doc: Document, latex: str, numbered: bool = True) -> None:
        """添加块级公式（公式居中，序号靠右）- 使用制表位实现"""
        para = doc.add_paragraph()
        
        # 设置制表位：居中制表位在页面中间，右对齐制表位在右边
        # A4页面可用宽度约6.27英寸
        pPr = para._p.get_or_add_pPr()
        tabs = OxmlElement('w:tabs')
        
        # 居中制表位 (页面中间 = 3.135英寸 ≈ 4478 twips，1英寸=1440 twips)
        tab_center = OxmlElement('w:tab')
        tab_center.set(qn('w:val'), 'center')
        tab_center.set(qn('w:pos'), '4478')
        tabs.append(tab_center)
        
        # 右对齐制表位 (页面右边 = 6.27英寸 ≈ 9029 twips)
        tab_right = OxmlElement('w:tab')
        tab_right.set(qn('w:val'), 'right')
        tab_right.set(qn('w:pos'), '9029')
        tabs.append(tab_right)
        
        pPr.append(tabs)
        
        # 添加Tab到中间，插入公式
        para.add_run('\t')
        try:
            self._insert_math(para, latex)
        except (ValueError, etree.XMLSyntaxError):
            run = para.add_run(latex)
            run.font.italic = True
        
        # 添加Tab到右边，插入序号
        if numbered:
            self.equation_counter += 1
            para.add_run('\t')
            run = para.add_run(f"({self.equation_counter})")
            run.font.size = FONT_SIZES['xiao_si']
    
    def add_inline_equation(self, paragraph, latex: str) -> None:
        """添加行内公式"""
        try:
            self._insert_math(paragraph, latex)
        except (ValueError, etree.XMLSyntaxError):
            # 转换失败时显示原始LaTeX（静默处理）
            run = paragraph.add_run(latex)
            run.font.italic = True
    
    def _insert_math(self, paragraph, latex: str) -> None:
        """将LaTeX公式插入到段落中"""
        # LaTeX转MathML
        mathml = latex_to_mathml(latex)
        
        # MathML转OMML (Office Math Markup Language)
        omml = self._mathml_to_omml(mathml)
        
        if omml is not None:
            # 插入到段落
            paragraph._p.append(omml)
        else:
            raise ValueError("OMML转换失败")
    
    def _mathml_to_omml(self, mathml: str):
        """MathML转换为OMML"""
        try:
            # 创建oMath元素
            nsmap_math = {'m': 'http://schemas.openxmlformats.org/officeDocument/2006/math'}
            oMath = etree.Element('{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath', 
                                  nsmap=nsmap_math)
            
            # 使用recover模式解析，容忍一些错误
            parser = etree.XMLParser(recover=True, encoding='utf-8')
            mathml_tree = etree.fromstring(mathml.encode('utf-8'), parser=parser)
            self._convert_mathml_element(mathml_tree, oMath)
            
            return oMath
            
        except etree.XMLSyntaxError:
            # 静默处理OMML转换错误
            return None
    
    def _decode_text(self, text: str) -> str:
        """解码HTML实体"""
        if not text:
            return ''
        import html
        return html.unescape(text)
    
    def _convert_mathml_element(self, mathml_elem, parent):
        """递归转换MathML元素到OMML"""
        ns = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
        
        tag = mathml_elem.tag.split('}')[-1] if '}' in mathml_elem.tag else mathml_elem.tag
        
        if tag in ('math', 'mrow'):
            # 容器元素，处理子元素
            for child in mathml_elem:
                self._convert_mathml_element(child, parent)
                
        elif tag == 'mi':
            # 标识符（变量）
            r = etree.SubElement(parent, f'{{{ns}}}r')
            t = etree.SubElement(r, f'{{{ns}}}t')
            t.text = self._decode_text(mathml_elem.text)
            
        elif tag == 'mn':
            # 数字
            r = etree.SubElement(parent, f'{{{ns}}}r')
            t = etree.SubElement(r, f'{{{ns}}}t')
            t.text = self._decode_text(mathml_elem.text)
            
        elif tag == 'mo':
            # 运算符
            r = etree.SubElement(parent, f'{{{ns}}}r')
            t = etree.SubElement(r, f'{{{ns}}}t')
            t.text = self._decode_text(mathml_elem.text)
            
        elif tag == 'mfrac':
            # 分数
            f = etree.SubElement(parent, f'{{{ns}}}f')
            num = etree.SubElement(f, f'{{{ns}}}num')
            den = etree.SubElement(f, f'{{{ns}}}den')
            
            children = list(mathml_elem)
            if len(children) >= 1:
                self._convert_mathml_element(children[0], num)
            if len(children) >= 2:
                self._convert_mathml_element(children[1], den)
                
        elif tag == 'msup':
            # 上标
            sSup = etree.SubElement(parent, f'{{{ns}}}sSup')
            e = etree.SubElement(sSup, f'{{{ns}}}e')
            sup = etree.SubElement(sSup, f'{{{ns}}}sup')
            
            children = list(mathml_elem)
            if len(children) >= 1:
                self._convert_mathml_element(children[0], e)
            if len(children) >= 2:
                self._convert_mathml_element(children[1], sup)
                
        elif tag == 'msub':
            # 下标
            sSub = etree.SubElement(parent, f'{{{ns}}}sSub')
            e = etree.SubElement(sSub, f'{{{ns}}}e')
            sub = etree.SubElement(sSub, f'{{{ns}}}sub')
            
            children = list(mathml_elem)
            if len(children) >= 1:
                self._convert_mathml_element(children[0], e)
            if len(children) >= 2:
                self._convert_mathml_element(children[1], sub)
                
        elif tag == 'msqrt':
            # 平方根
            rad = etree.SubElement(parent, f'{{{ns}}}rad')
            radPr = etree.SubElement(rad, f'{{{ns}}}radPr')
            degHide = etree.SubElement(radPr, f'{{{ns}}}degHide')
            degHide.set(f'{{{ns}}}val', '1')
            e = etree.SubElement(rad, f'{{{ns}}}e')
            
            for child in mathml_elem:
                self._convert_mathml_element(child, e)
                
        elif tag == 'msubsup':
            # 上下标
            sSubSup = etree.SubElement(parent, f'{{{ns}}}sSubSup')
            e = etree.SubElement(sSubSup, f'{{{ns}}}e')
            sub = etree.SubElement(sSubSup, f'{{{ns}}}sub')
            sup = etree.SubElement(sSubSup, f'{{{ns}}}sup')
            
            children = list(mathml_elem)
            if len(children) >= 1:
                self._convert_mathml_element(children[0], e)
            if len(children) >= 2:
                self._convert_mathml_element(children[1], sub)
            if len(children) >= 3:
                self._convert_mathml_element(children[2], sup)
        
        elif tag == 'mtable':
            # 表格/矩阵（用于aligned环境）
            m = etree.SubElement(parent, f'{{{ns}}}m')
            mPr = etree.SubElement(m, f'{{{ns}}}mPr')
            
            # 处理每一行
            for row_elem in mathml_elem:
                mr = etree.SubElement(m, f'{{{ns}}}mr')
                row_tag = row_elem.tag.split('}')[-1] if '}' in row_elem.tag else row_elem.tag
                if row_tag == 'mtr':
                    for cell_elem in row_elem:
                        cell_tag = cell_elem.tag.split('}')[-1] if '}' in cell_elem.tag else cell_elem.tag
                        if cell_tag == 'mtd':
                            e = etree.SubElement(mr, f'{{{ns}}}e')
                            for child in cell_elem:
                                self._convert_mathml_element(child, e)
        
        elif tag == 'mover':
            # 上划线/帽子
            acc = etree.SubElement(parent, f'{{{ns}}}acc')
            accPr = etree.SubElement(acc, f'{{{ns}}}accPr')
            e = etree.SubElement(acc, f'{{{ns}}}e')
            
            children = list(mathml_elem)
            if len(children) >= 1:
                self._convert_mathml_element(children[0], e)
            # 处理上标符号
            if len(children) >= 2 and children[1].text:
                chr_elem = etree.SubElement(accPr, f'{{{ns}}}chr')
                chr_elem.set(f'{{{ns}}}val', children[1].text)
        
        elif tag == 'munder':
            # 下标记
            for child in mathml_elem:
                self._convert_mathml_element(child, parent)
        
        elif tag == 'munderover':
            # 上下标记（如求和符号）
            nary = etree.SubElement(parent, f'{{{ns}}}nary')
            naryPr = etree.SubElement(nary, f'{{{ns}}}naryPr')
            sub = etree.SubElement(nary, f'{{{ns}}}sub')
            sup = etree.SubElement(nary, f'{{{ns}}}sup')
            e = etree.SubElement(nary, f'{{{ns}}}e')
            
            children = list(mathml_elem)
            if len(children) >= 1:
                # 第一个是操作符
                if children[0].text:
                    chr_elem = etree.SubElement(naryPr, f'{{{ns}}}chr')
                    chr_elem.set(f'{{{ns}}}val', self._decode_text(children[0].text))
            if len(children) >= 2:
                self._convert_mathml_element(children[1], sub)
            if len(children) >= 3:
                self._convert_mathml_element(children[2], sup)
        
        elif tag == 'mtext':
            # 文本
            r = etree.SubElement(parent, f'{{{ns}}}r')
            rPr = etree.SubElement(r, f'{{{ns}}}rPr')
            nor = etree.SubElement(rPr, f'{{{ns}}}nor')  # 正体
            t = etree.SubElement(r, f'{{{ns}}}t')
            t.text = self._decode_text(mathml_elem.text)
        
        elif tag == 'mspace':
            # 空格
            r = etree.SubElement(parent, f'{{{ns}}}r')
            t = etree.SubElement(r, f'{{{ns}}}t')
            t.text = ' '
        
        elif tag == 'mfenced':
            # 括号包围
            d = etree.SubElement(parent, f'{{{ns}}}d')
            dPr = etree.SubElement(d, f'{{{ns}}}dPr')
            
            # 设置括号类型
            open_char = mathml_elem.get('open', '(')
            close_char = mathml_elem.get('close', ')')
            begChr = etree.SubElement(dPr, f'{{{ns}}}begChr')
            begChr.set(f'{{{ns}}}val', open_char)
            endChr = etree.SubElement(dPr, f'{{{ns}}}endChr')
            endChr.set(f'{{{ns}}}val', close_char)
            
            e = etree.SubElement(d, f'{{{ns}}}e')
            for child in mathml_elem:
                self._convert_mathml_element(child, e)
        
        else:
            # 其他元素，尝试处理子元素
            for child in mathml_elem:
                self._convert_mathml_element(child, parent)
            if mathml_elem.text:
                r = etree.SubElement(parent, f'{{{ns}}}r')
                t = etree.SubElement(r, f'{{{ns}}}t')
                t.text = self._decode_text(mathml_elem.text)


class CodeHandler:
    """代码块处理器"""
    
    def add_code_block(self, doc: Document, code: str, language: str = None) -> None:
        """添加代码块"""
        # 添加语言标识（如果有）
        if language:
            lang_para = doc.add_paragraph()
            lang_run = lang_para.add_run(f"代码 ({language}):")
            lang_run.font.size = FONT_SIZES['caption']
            lang_run.font.bold = True
        
        # 添加代码内容
        lines = code.split('\n')
        for line in lines:
            para = doc.add_paragraph(style='Code Block')
            run = para.add_run(line)
            run.font.name = FONTS['mono']
            run.font.size = FONT_SIZES['code']
            
            # 设置背景色
            self._set_paragraph_shading(para, 'F5F5F5')
    
    def add_inline_code(self, paragraph, code: str) -> None:
        """添加行内代码"""
        run = paragraph.add_run(code)
        run.font.name = FONTS['mono']
        run.font.size = FONT_SIZES['code']
        
        # 设置背景
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'F0F0F0')
        run._r.get_or_add_rPr().append(shading)
    
    def _set_paragraph_shading(self, paragraph, color: str) -> None:
        """设置段落背景色"""
        pPr = paragraph._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), color)
        pPr.append(shd)


class ListHandler:
    """列表处理器"""
    
    def add_bullet_item(self, doc: Document, text: str, level: int = 0) -> None:
        """添加无序列表项"""
        para = doc.add_paragraph(style='List Bullet')
        para.add_run(text)
        
        # 设置缩进级别
        if level > 0:
            para.paragraph_format.left_indent = Cm(1.27 * level)
    
    def add_numbered_item(self, doc: Document, text: str, level: int = 0) -> None:
        """添加有序列表项"""
        para = doc.add_paragraph(style='List Number')
        para.add_run(text)
        
        if level > 0:
            para.paragraph_format.left_indent = Cm(1.27 * level)
