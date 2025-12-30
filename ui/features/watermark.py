# -*- coding: utf-8 -*-
"""
水印功能模块
"""

import customtkinter as ctk
from tkinter import filedialog, messagebox, colorchooser
from typing import TYPE_CHECKING, Optional, Dict
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

if TYPE_CHECKING:
    from gui import App


class WatermarkFeature:
    """Word文档水印功能"""
    
    def __init__(self, app: 'App'):
        self.app = app
        self.watermark_config = None
    
    def show_watermark_dialog(self):
        """显示水印配置对话框"""
        dialog = WatermarkDialog(self.app, self)
        dialog.grab_set()
        dialog.wait_window()
    
    def apply_watermark_to_docx(self, doc: Document, config: Dict):
        """
        应用水印到Word文档
        
        Args:
            doc: python-docx Document对象
            config: 水印配置
                {
                    'type': 'text' | 'image',
                    'text': str,  # 文字水印
                    'color': (r, g, b),
                    'opacity': float,  # 0.0 - 1.0
                    'angle': int,  # 旋转角度
                    'font_size': int,
                    'image_path': str,  # 图片水印
                }
        """
        try:
            if config['type'] == 'text':
                self._add_text_watermark(doc, config)
            elif config['type'] == 'image':
                self._add_image_watermark(doc, config)
            
            return True
        except Exception as e:
            print(f"应用水印失败: {e}")
            return False
    
    def _add_text_watermark(self, doc: Document, config: Dict):
        """添加文字水印"""
        # Word水印是通过在页眉中添加文本框实现的
        section = doc.sections[0]
        header = section.header
        
        # 清除现有水印
        for paragraph in header.paragraphs:
            paragraph.clear()
        
        # 创建水印段落
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        
        # 使用WordprocessingML添加水印
        # 注意：python-docx对水印的支持有限，这里使用底层XML
        watermark_xml = self._generate_text_watermark_xml(config)
        paragraph._element.getparent().insert(0, parse_xml(watermark_xml))
    
    def _generate_text_watermark_xml(self, config: Dict) -> str:
        """生成文字水印的XML"""
        text = config.get('text', 'DRAFT')
        color_rgb = config.get('color', (192, 192, 192))
        opacity = config.get('opacity', 0.5)
        angle = config.get('angle', -45)
        font_size = config.get('font_size', 72)
        
        # 将RGB颜色转换为十六进制
        color_hex = '{:02X}{:02X}{:02X}'.format(*color_rgb)
        
        # 计算透明度（0-1 转换为 0-100000）
        opacity_val = int(opacity * 100000)
        
        # 生成VML水印
        xml = f'''
        <w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
             xmlns:v="urn:schemas-microsoft-com:vml"
             xmlns:o="urn:schemas-microsoft-com:office:office">
            <w:pPr>
                <w:pStyle w:val="Header"/>
            </w:pPr>
            <w:r>
                <w:rPr/>
                <w:pict>
                    <v:shapetype id="_x0000_t136" coordsize="21600,21600" o:spt="136" adj="10800" path="m@7,l@8,m@5,21600l@6,21600e">
                        <v:f eqn="sum #0 0 10800"/>
                        <v:f eqn="prod #0 2 1"/>
                        <v:f eqn="sum 21600 0 @1"/>
                        <v:f eqn="sum 0 0 @2"/>
                        <v:f eqn="sum 21600 0 @3"/>
                        <v:f eqn="if @0 @3 0"/>
                        <v:f eqn="if @0 21600 @1"/>
                        <v:f eqn="if @0 0 @2"/>
                        <v:f eqn="if @0 @4 21600"/>
                        <v:f eqn="mid @5 @6"/>
                        <v:f eqn="mid @8 @5"/>
                        <v:f eqn="mid @7 @8"/>
                        <v:f eqn="mid @6 @7"/>
                        <v:f eqn="sum @6 0 @5"/>
                        <v:path textpathok="t" o:connecttype="custom" o:connectlocs="@9,0;@10,10800;@11,21600;@12,10800" o:connectangles="270,180,90,0"/>
                        <v:textpath on="t" fitshape="t"/>
                        <v:handles>
                            <v:h position="#0,bottomRight" xrange="6629,14971"/>
                        </v:handles>
                        <o:lock v:ext="edit" text="t" shapetype="t"/>
                    </v:shapetype>
                    <v:shape id="PowerPlusWaterMarkObject" o:spid="_x0000_s2050" type="#_x0000_t136" 
                             style="position:absolute;margin-left:0;margin-top:0;width:467.75pt;height:467.75pt;rotation:{angle};z-index:-251658240;
                             mso-position-horizontal:center;mso-position-horizontal-relative:margin;
                             mso-position-vertical:center;mso-position-vertical-relative:margin" 
                             o:allowincell="f" fillcolor="#{color_hex}" stroked="f">
                        <v:fill opacity="{opacity_val}f"/>
                        <v:textpath style="font-family:&quot;Calibri&quot;;font-size:{font_size}pt" string="{text}"/>
                    </v:shape>
                </w:pict>
            </w:r>
        </w:p>
        '''
        
        return xml.strip()
    
    def _add_image_watermark(self, doc: Document, config: Dict):
        """添加图片水印"""
        # 图片水印实现比较复杂，需要在页眉中添加图片并设置透明度
        # 这里提供一个简化版本
        try:
            section = doc.sections[0]
            header = section.header
            
            # 清除现有内容
            for paragraph in header.paragraphs:
                paragraph.clear()
            
            # 添加图片
            paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            run = paragraph.add_run()
            
            image_path = config.get('image_path')
            if image_path:
                # 添加图片（简化版，不支持透明度和旋转）
                run.add_picture(image_path, width=Cm(10))
                
                # 居中
                paragraph.alignment = 1  # CENTER
                
        except Exception as e:
            print(f"添加图片水印失败: {e}")
            raise


class WatermarkDialog(ctk.CTkToplevel):
    """水印配置对话框"""
    
    def __init__(self, parent, watermark_feature: WatermarkFeature):
        super().__init__(parent)
        
        self.watermark_feature = watermark_feature
        self.color_rgb = (192, 192, 192)  # 默认灰色
        
        # 窗口配置
        self.title("水印设置")
        self.geometry("450x500")
        self.resizable(False, False)
        
        # 居中显示
        self.update_idletasks()
        x = (self.winfo_screenwidth() // 2) - (450 // 2)
        y = (self.winfo_screenheight() // 2) - (500 // 2)
        self.geometry(f"+{x}+{y}")
        
        self._create_ui()
    
    def _create_ui(self):
        """创建界面"""
        # 标题
        title_label = ctk.CTkLabel(
            self,
            text="💧 水印设置",
            font=ctk.CTkFont(size=18, weight="bold")
        )
        title_label.pack(pady=20)
        
        # 水印类型选择
        type_frame = ctk.CTkFrame(self, fg_color="transparent")
        type_frame.pack(pady=10, padx=20, fill="x")
        
        ctk.CTkLabel(
            type_frame,
            text="水印类型：",
            font=ctk.CTkFont(size=14)
        ).pack(side="left", padx=10)
        
        self.type_var = ctk.StringVar(value="text")
        
        ctk.CTkRadioButton(
            type_frame,
            text="文字水印",
            variable=self.type_var,
            value="text",
            command=self._on_type_change
        ).pack(side="left", padx=10)
        
        ctk.CTkRadioButton(
            type_frame,
            text="图片水印",
            variable=self.type_var,
            value="image",
            command=self._on_type_change
        ).pack(side="left", padx=10)
        
        # 配置容器
        self.config_container = ctk.CTkFrame(self)
        self.config_container.pack(pady=10, padx=20, fill="both", expand=True)
        
        # 显示文字水印配置
        self._create_text_watermark_ui()
    
    def _on_type_change(self):
        """水印类型改变"""
        # 清除现有配置UI
        for widget in self.config_container.winfo_children():
            widget.destroy()
        
        # 根据类型显示不同的配置
        if self.type_var.get() == "text":
            self._create_text_watermark_ui()
        else:
            self._create_image_watermark_ui()
    
    def _create_text_watermark_ui(self):
        """创建文字水印配置UI"""
        # 水印文本
        text_frame = ctk.CTkFrame(self.config_container, fg_color="transparent")
        text_frame.pack(pady=10, padx=10, fill="x")
        
        ctk.CTkLabel(text_frame, text="水印文本：").pack(anchor="w")
        self.text_entry = ctk.CTkEntry(text_frame, width=360)
        self.text_entry.pack(pady=5)
        self.text_entry.insert(0, "DRAFT")
        
        # 字体大小
        size_frame = ctk.CTkFrame(self.config_container, fg_color="transparent")
        size_frame.pack(pady=10, padx=10, fill="x")
        
        ctk.CTkLabel(size_frame, text="字体大小：").pack(side="left")
        self.font_size_var = ctk.IntVar(value=72)
        ctk.CTkSlider(
            size_frame,
            from_=24,
            to=144,
            variable=self.font_size_var,
            width=250
        ).pack(side="left", padx=10)
        self.font_size_label = ctk.CTkLabel(size_frame, text="72")
        self.font_size_label.pack(side="left")
        self.font_size_var.trace_add('write', self._update_font_size_label)
        
        # 颜色
        color_frame = ctk.CTkFrame(self.config_container, fg_color="transparent")
        color_frame.pack(pady=10, padx=10, fill="x")
        
        ctk.CTkLabel(color_frame, text="颜色：").pack(side="left")
        self.color_button = ctk.CTkButton(
            color_frame,
            text="选择颜色",
            command=self._choose_color,
            width=100
        )
        self.color_button.pack(side="left", padx=10)
        self.color_preview = ctk.CTkLabel(
            color_frame,
            text="      ",
            fg_color="#C0C0C0",
            corner_radius=5,
            width=50
        )
        self.color_preview.pack(side="left", padx=5)
        
        # 透明度
        opacity_frame = ctk.CTkFrame(self.config_container, fg_color="transparent")
        opacity_frame.pack(pady=10, padx=10, fill="x")
        
        ctk.CTkLabel(opacity_frame, text="透明度：").pack(side="left")
        self.opacity_var = ctk.DoubleVar(value=0.5)
        ctk.CTkSlider(
            opacity_frame,
            from_=0.1,
            to=1.0,
            variable=self.opacity_var,
            width=250
        ).pack(side="left", padx=10)
        self.opacity_label = ctk.CTkLabel(opacity_frame, text="50%")
        self.opacity_label.pack(side="left")
        self.opacity_var.trace_add('write', self._update_opacity_label)
        
        # 旋转角度
        angle_frame = ctk.CTkFrame(self.config_container, fg_color="transparent")
        angle_frame.pack(pady=10, padx=10, fill="x")
        
        ctk.CTkLabel(angle_frame, text="旋转角度：").pack(side="left")
        self.angle_var = ctk.IntVar(value=-45)
        ctk.CTkSlider(
            angle_frame,
            from_=-90,
            to=90,
            variable=self.angle_var,
            width=250
        ).pack(side="left", padx=10)
        self.angle_label = ctk.CTkLabel(angle_frame, text="-45°")
        self.angle_label.pack(side="left")
        self.angle_var.trace_add('write', self._update_angle_label)
        
        # 按钮
        self._create_buttons()
    
    def _create_image_watermark_ui(self):
        """创建图片水印配置UI"""
        # 图片选择
        image_frame = ctk.CTkFrame(self.config_container, fg_color="transparent")
        image_frame.pack(pady=20, padx=10, fill="x")
        
        ctk.CTkLabel(image_frame, text="选择图片：").pack(anchor="w")
        
        select_frame = ctk.CTkFrame(image_frame, fg_color="transparent")
        select_frame.pack(pady=10, fill="x")
        
        self.image_path_var = ctk.StringVar(value="")
        self.image_entry = ctk.CTkEntry(
            select_frame,
            textvariable=self.image_path_var,
            width=260
        )
        self.image_entry.pack(side="left", padx=5)
        
        ctk.CTkButton(
            select_frame,
            text="浏览...",
            command=self._browse_image,
            width=80
        ).pack(side="left")
        
        # 提示
        ctk.CTkLabel(
            self.config_container,
            text="💡 提示：建议使用PNG格式的透明背景图片",
            font=ctk.CTkFont(size=11),
            text_color="gray"
        ).pack(pady=20)
        
        # 按钮
        self._create_buttons()
    
    def _create_buttons(self):
        """创建底部按钮"""
        button_frame = ctk.CTkFrame(self, fg_color="transparent")
        button_frame.pack(pady=15, padx=20, fill="x", side="bottom")
        
        ctk.CTkButton(
            button_frame,
            text="应用水印",
            command=self._apply_watermark,
            width=150
        ).pack(side="left", padx=5)
        
        ctk.CTkButton(
            button_frame,
            text="取消",
            command=self.destroy,
            fg_color="gray",
            width=150
        ).pack(side="right", padx=5)
    
    def _choose_color(self):
        """选择颜色"""
        color = colorchooser.askcolor(title="选择水印颜色")
        if color[0]:  # color[0] is RGB tuple
            self.color_rgb = tuple(int(c) for c in color[0])
            # 更新预览
            color_hex = '#{:02X}{:02X}{:02X}'.format(*self.color_rgb)
            self.color_preview.configure(fg_color=color_hex)
    
    def _browse_image(self):
        """浏览图片"""
        file_path = filedialog.askopenfilename(
            title="选择水印图片",
            filetypes=[
                ("图片文件", "*.png *.jpg *.jpeg *.bmp"),
                ("所有文件", "*.*")
            ]
        )
        if file_path:
            self.image_path_var.set(file_path)
    
    def _update_font_size_label(self, *args):
        """更新字体大小标签"""
        self.font_size_label.configure(text=str(self.font_size_var.get()))
    
    def _update_opacity_label(self, *args):
        """更新透明度标签"""
        opacity = self.opacity_var.get()
        self.opacity_label.configure(text=f"{int(opacity * 100)}%")
    
    def _update_angle_label(self, *args):
        """更新角度标签"""
        self.angle_label.configure(text=f"{self.angle_var.get()}°")
    
    def _apply_watermark(self):
        """应用水印"""
        watermark_type = self.type_var.get()
        
        if watermark_type == "text":
            config = {
                'type': 'text',
                'text': self.text_entry.get().strip(),
                'color': self.color_rgb,
                'opacity': self.opacity_var.get(),
                'angle': self.angle_var.get(),
                'font_size': self.font_size_var.get(),
            }
            
            if not config['text']:
                messagebox.showwarning("警告", "请输入水印文本！")
                return
        else:
            config = {
                'type': 'image',
                'image_path': self.image_path_var.get(),
            }
            
            if not config['image_path']:
                messagebox.showwarning("警告", "请选择水印图片！")
                return
        
        # 保存配置
        self.watermark_feature.watermark_config = config
        
        messagebox.showinfo(
            "成功",
            "水印配置已保存！\n\n导出Word文档时将自动应用水印。"
        )
        
        self.destroy()
